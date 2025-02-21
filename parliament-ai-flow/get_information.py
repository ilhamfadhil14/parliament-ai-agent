import requests
import json
import os

from dotenv import load_dotenv
from datetime import datetime
from promptflow.core import tool
from openai import OpenAI
from typing import List, Dict, Any
from duckduckgo_search import DDGS
from docx import Document
from docx.shared import Inches

load_dotenv()

EMBEDDING_MODEL_NAME = os.environ.get('EMBEDDING_MODEL_NAME')
SEARCH_SERVICE_NAME = os.environ.get('SEARCH_SERVICE_NAME')
INDEX_NAME = os.environ.get('INDEX_NAME')
SEARCH_API_KEY = os.environ.get('SEARCH_API_KEY')
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
ACCESS_TOKEN = os.environ.get('ACCESS_TOKEN')

client = OpenAI(api_key=OPENAI_API_KEY)
search_cliet = DDGS(verify=False)

def get_latest_news(query:str) -> list:
    
    response = search_cliet.news(query, max_results=5)

    return response

def create_embeddings(prompt:str) -> list:

    response = client.embeddings.create(
        model=EMBEDDING_MODEL_NAME,
        input=prompt
    )

    return response.data[0].embedding

def get_policy_database(
    prompt:str
    ) -> Dict[Any, Any]:

    embedding = create_embeddings(prompt)
    result = search_documents(search_service_name=SEARCH_SERVICE_NAME, 
                              index_name=INDEX_NAME, 
                              api_key=SEARCH_API_KEY, 
                              vector=embedding)

    return result

def search_documents(
    search_service_name: str,
    index_name: str,
    api_key: str,
    vector: List,
    select_fields: str = "title, chunk",
    top_k: int = 5,
    api_version: str = "2024-07-01"
) -> Dict[Any, Any]:
    """
    Perform vector search against Azure Cognitive Search index
    
    Args:
        search_service_name: Name of your search service
        index_name: Name of the search index
        api_key: Admin API key for authentication
        vector: List of vector embeddings
        select_fields: Comma-separated fields to return
        top_k: Number of results to return
    """
    
    # Construct the endpoint URL
    endpoint = f"https://{search_service_name}.search.windows.net"
    url = f"{endpoint}/indexes/{index_name}/docs/search?api-version={api_version}"
    
    # Set up headers
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key
    }
    
    # Construct request body
    body = {
        "count": True,
        "select": select_fields,
        "vectorQueries": [
            {
                "kind": "vector",
                "vector": vector,
                "exhaustive": True,
                "fields": "text_vector",
                "weight": 0.5,
                "k": top_k
            }
        ]
    }
    
    try:
        # Make POST request
        response = requests.post(url, headers=headers, json=body)
        response.raise_for_status()
        
        results = response.json()
        filtered_results = [{'chunk': res['chunk'], 'title': res['title']} for res in results['value']]
        return filtered_results
    except requests.exceptions.RequestException as e:
        raise Exception(f"Search request failed: {str(e)}")

def write_to_microsoft_word(content: str, title: str = "Generated Report") -> str:
    """
    Write content to a Microsoft Word document.
    
    Args:
        content: The content to write to the document
        title: The title of the document (default: "Generated Report")
        
    Returns:
        str: Path to the saved document
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {timestamp}")
        
        # Add content
        doc.add_paragraph(content)
        
        # Create output directory if it doesn't exist
        output_dir = "output"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # Generate filename based on title and timestamp
        filename = f"{title.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(output_dir, filename)
        
        # Save the document
        doc.save(file_path)
        
        return file_path
        
    except Exception as e:
        raise Exception(f"Failed to write to Word document: {str(e)}")

def create_email_draft(content:str, title: str = "Generated Email Draft") -> str:
    """
    Create an email draft with the specified content.
    
    Args:
        content: The content to include in the email
        title: The title of the email (default: "Generated Email Draft")
        
    Returns:
        str: The email draft content
    """

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    request_body = {
    "subject":f"{title} - {timestamp}",
    "importance":"Low",
    "body":{
        "contentType":"HTML",
        "content": content
        }
    }

    # API endpoint for creating draft email
    endpoint = "https://graph.microsoft.com/v1.0/me/messages"

    # Headers with auth token
    headers = {
        "Authorization": f"Bearer {ACCESS_TOKEN}",
        "Content-Type": "application/json"
    }

    # Send POST request
    response = requests.post(endpoint, headers=headers, json=request_body)
    response.raise_for_status()

    return response.json()

@tool
def run_function(response_message:dict) -> str:
    function_call = response_message.get('function_call', None)

    if function_call and "name" in function_call and "arguments" in function_call:
        function_name = function_call["name"]
        function_args = json.loads(function_call["arguments"])
        print(function_args)

        result = globals()[function_name](**function_args)

    else:
        print("No function call found in response message")
        if isinstance(response_message, dict):
            result = response_message.get("content", "")
        else:
            result = response_message
    return result